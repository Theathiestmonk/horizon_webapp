import logging
import os
import shutil
import subprocess
import tempfile
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict
from docx import Document as DocxDocument
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Length, Pt
from docxtpl import DocxTemplate
from num2words import num2words

logger = logging.getLogger(__name__)

TEMPLATE_DIR = Path("templates")
OUTPUT_DIR = Path("output")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
BASE_DOWNLOAD_URL = os.getenv("BASE_DOWNLOAD_URL", "http://localhost:8080/output").rstrip("/")

# ── One-page enforcement ─────────────────────────────────────────────────────
# Every generated document except DBK MUST render to exactly one page —
# client requirement, no exceptions. Starting at ONE_PAGE_DEFAULT_PT, font
# size is dynamically searched downward in ONE_PAGE_STEP_PT increments (not a
# handful of fixed named sizes) until the document actually fits, all the way
# down to ONE_PAGE_MIN_PT if that's what it genuinely takes — unlike the
# fixed 5-step ladder this replaces, there is no per-template floor that can
# leave a heavy shipment stuck on 2 pages. ONE_PAGE_MIN_PT is deliberately
# still comfortably legible (6pt is small but readable), not a token
# stopping point — going lower than that stops helping in practice (row/cell
# padding and fixed-size elements like the logo start to dominate the page
# long before individual characters would become unreadable), so a document
# that still doesn't fit at 6pt is accepted as a genuine, exceptional
# overflow rather than pushed into truly illegible territory.
#
# "Fits" is verified for real (not guessed) by round-tripping through
# LibreOffice headless to PDF and counting actual pages on every single step
# — python-docx has no layout engine of its own, so there is no way to know
# how many pages a .docx will render to without actually rendering it.
ONE_PAGE_EXCLUDED_TEMPLATES = {"DBK_Declaration.docx"}
ONE_PAGE_DEFAULT_PT = 11.0
ONE_PAGE_MIN_PT = 6.0
ONE_PAGE_STEP_PT = 0.5
ONE_PAGE_SOFFICE_TIMEOUT_SEC = 60


def _one_page_size_ladder(ceiling_pt: float, floor_pt: float) -> list:
    """Every font size to try, largest first, from ceiling_pt down to
    floor_pt in ONE_PAGE_STEP_PT increments — the dynamic replacement for a
    fixed handful of named sizes. Always includes floor_pt itself even if it
    doesn't land exactly on a step boundary, so the true floor is always the
    last-resort attempt regardless of the ceiling/step combination."""
    sizes = []
    size = ceiling_pt
    while size > floor_pt:
        sizes.append(round(size, 1))
        size -= ONE_PAGE_STEP_PT
    sizes.append(round(floor_pt, 1))
    return sizes

# ANX C has enough real content (Annexure C's 15 fixed line items, each with
# its own row) that 11pt/10.5pt almost always overflow to a 2nd page anyway —
# confirmed by reproduction. Starting its search at 10pt instead just skips
# those two known-futile attempts (each one costs a real LibreOffice
# round-trip, ~2s) — a pure performance optimization, not a correctness
# floor/ceiling; it still searches all the way down to ONE_PAGE_MIN_PT like
# every other template if it ever needs to.
ONE_PAGE_TEMPLATE_MAX_PT_OVERRIDE = {
    "ANX C.docx": 10.0,
}

# Applied to every "at least" row height on EVERY ladder step, including the
# very first (at ONE_PAGE_DEFAULT_PT) — without this, row-height scaling
# only kicked in once the ladder had already dropped BELOW the default font
# size (ratio was exactly 1.0 on the first try), so genuinely wasted
# template-authored padding never got reclaimed on the one attempt that
# matters most: whether the document fits at the default size at all. Empirically
# confirmed via a real reproduction (PI FORMAT with 5 vehicles, which spilled
# one line onto an otherwise-empty page 2 purely from this unclaimed padding):
# 0.85 was the threshold that brought it back to one page at the full
# ONE_PAGE_DEFAULT_PT, with no visible clipping — "at least" rows can never
# be clipped by this, they simply grow back past the reduced minimum if
# their actual content needs more room.
ONE_PAGE_ROW_HEIGHT_BASELINE_RATIO = 0.85

# Each document's own title heading (e.g. "COMMERCIAL INVOICE", "PACKING
# LIST") is exempted from the uniform shrink below and pinned to this size
# instead — the client wants the document name to stay legible even at the
# smallest compression step, while every other run still follows the ladder.
# Matched by exact paragraph text (post-trim, case-insensitive) against the
# literal title text baked into each template — confirmed via inspection of
# templates/*.docx (some, like "ANX C.docx", carry the title inside a table
# cell rather than a body paragraph; _iter_all_paragraphs already covers both).
ONE_PAGE_TITLE_PT = 17
ONE_PAGE_TITLE_TEXTS = {
    "COMMERCIAL INVOICE", "PACKING LIST", "TAX INVOICE", "PROFORMA INVOICE",
    "ANNEXURE -1", "ANNEXURE C",
}

# The exporter's static "HORIZON ENTERPRISE" name + address block that sits
# at the bottom of most templates (below the main content, outside any
# {{ exporter.* }} placeholder) — same pin-instead-of-shrink treatment as
# the title above. Company name matched by exact text; the address line
# varies slightly per template (with/without "Email:.../Website:...", "E-802"
# vs "E 802" spacing — confirmed via inspection of templates/*.docx) so it's
# matched by prefix/substring instead of an exact-text set.
ONE_PAGE_COMPANY_NAME_PT = 14
ONE_PAGE_COMPANY_NAME_TEXTS = {"HORIZON ENTERPRISE"}
ONE_PAGE_ADDRESS_PT = 12

# The "AMOUNT CHARGEABLE..." words-in-full line (e.g. CHA TI's "AMOUNT
# CHARGEABLE INR: Rupees Seventy Eight Lakh...Only") is long enough that the
# uniform ladder size (11pt default) wraps it onto two lines — the template
# already had this line hand-set to 10pt specifically so it fits on one line,
# but that got silently overwritten back to the uniform size on every
# generation since nothing exempted it from the shrink-everything-together
# pass. Pinned the same way as the title/company/address blocks above so it
# stays at 10pt regardless of what the rest of the page is doing.
ONE_PAGE_AMOUNT_CHARGEABLE_PT = 10


def _is_footer_address_paragraph(text_upper: str) -> bool:
    return (
        text_upper.startswith("E-802") or text_upper.startswith("E 802")
        or "SAFAL PARISAR" in text_upper
        or text_upper.startswith("EMAIL:") or text_upper.startswith("EMAIL :")
        or "WEBSITE:" in text_upper
    )


def _is_amount_chargeable_paragraph(text_upper: str) -> bool:
    return text_upper.startswith("AMOUNT CHARGEABLE")


def _iter_table_paragraphs(table):
    """Yields every paragraph in a table's cells, recursing into nested tables."""
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _iter_all_paragraphs(doc: "DocxDocument"):
    """Yields every paragraph in the document: body, all tables (incl. nested),
    and every section's header/footer (and any tables inside those)."""
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        yield from _iter_table_paragraphs(t)
    for section in doc.sections:
        for part in (section.header, section.footer):
            for p in part.paragraphs:
                yield p
            for t in part.tables:
                yield from _iter_table_paragraphs(t)


def _iter_body_and_header_footer_paragraphs(doc: "DocxDocument"):
    """Yields only NON-table paragraphs: plain body paragraphs plus each
    section's header/footer paragraphs — deliberately excluding every table
    cell (doc.tables AND any header/footer tables).

    Used only for the company-name/address pin below. Several templates
    render {{ exporter.company_name }} / {{ exporter.address }} more than
    once — the top "EXPORTER / MANUFACTURER" block and a "FOR
    {{ exporter.company_name }}" signature line both live in table cells —
    and once Jinja substitutes the real company name into both, they become
    text-identical to the genuine footer occurrence, so text matching alone
    can't tell them apart. Structural position can: inspection of
    templates/*.docx confirmed the real "HORIZON ENTERPRISE" + address block
    is always either a plain body paragraph (Packing_List.docx) or a true
    Word section footer (CHA TI.docx, CHA PL.docx) — never a table cell in
    either case — so restricting the pin to this iterator is what correctly
    excludes the in-table lookalikes.
    """
    for p in doc.paragraphs:
        yield p
    for section in doc.sections:
        for part in (section.header, section.footer):
            for p in part.paragraphs:
                yield p


def _iter_table_tables(table):
    """Yields a table and every table nested inside its cells, recursively."""
    yield table
    for row in table.rows:
        for cell in row.cells:
            for nested in cell.tables:
                yield from _iter_table_tables(nested)


def _iter_all_tables(doc: "DocxDocument"):
    """Yields every table in the document: body, every section's header/
    footer, and all nested tables — mirrors _iter_all_paragraphs above."""
    for t in doc.tables:
        yield from _iter_table_tables(t)
    for section in doc.sections:
        for part in (section.header, section.footer):
            for t in part.tables:
                yield from _iter_table_tables(t)


def _snapshot_row_heights(doc: "DocxDocument"):
    """Captures every table row's original (template-authored) 'at least'
    minimum height, once, before any ladder step touches it. Rows with
    hRule=EXACTLY or AUTO are left alone — EXACTLY clips overflow content if
    shrunk incorrectly, and AUTO already sizes itself to content. 'At least'
    rows are the only ones safe to shrink and the only kind found in the
    current templates (confirmed by inspection).

    Why this exists: shrinking font size alone (_set_all_font_sizes) doesn't
    reclaim page space held by static header/footer rows, because their
    minimum heights were authored around the template's original, larger
    font and don't shrink just because the text inside got smaller — the
    row simply gains extra empty padding instead. A row whose content
    genuinely needs more room (e.g. the per-vehicle description list) still
    grows past this minimum regardless, so scaling the minimum down never
    clips real content — it only removes the leftover padding that was
    sized for a bigger font.
    """
    snapshot = []
    for table in _iter_all_tables(doc):
        for row in table.rows:
            if row.height_rule == WD_ROW_HEIGHT_RULE.AT_LEAST and row.height:
                snapshot.append((row, int(row.height)))
    return snapshot


def _scale_row_heights(snapshot, ratio: float) -> None:
    """Rewrites every snapshotted row's minimum height to `ratio` of its
    original (template-authored) value — absolute, not cumulative, so this
    is safe to call again on every ladder step regardless of order."""
    for row, original_emu in snapshot:
        row.height = Length(max(1, round(original_emu * ratio)))


# Page margins are pure whitespace — unlike row height minimums (which
# protect real content that might need more room) or the pinned title/
# company/address sizes (which are a deliberate client requirement), there
# is nothing a shrunk top/bottom margin could ever clip, so this is safe to
# push much harder than the row-height ratio. Confirmed necessary by direct
# reproduction: a heavy Packing List (3 models, 8 vehicles, full chassis
# breakdown) still had the pinned "HORIZON ENTERPRISE" footer block spill a
# single line onto an otherwise-blank 2nd page even at the font-size floor
# with row heights fully reclaimed — shrinking top/bottom margins closed
# nearly all of that gap on its own. ONE_PAGE_MARGIN_FLOOR_PT keeps a small
# amount of real margin rather than going to 0 (which some renderers clamp
# to their own minimum anyway, making a literal 0 no more effective than a
# small positive value while looking broken if a viewer doesn't clamp it).
ONE_PAGE_MARGIN_FLOOR_PT = 10.0


def _snapshot_margins(doc: "DocxDocument"):
    """Captures every section's original top/bottom margin once, before any
    ladder step touches it — mirrors _snapshot_row_heights above."""
    return [(section, section.top_margin, section.bottom_margin) for section in doc.sections]


def _scale_margins(snapshot, ratio: float) -> None:
    """Rewrites every snapshotted section's top/bottom margin to `ratio` of
    its original value, never going below ONE_PAGE_MARGIN_FLOOR_PT — absolute,
    not cumulative, so safe to call again on every ladder step."""
    floor = Pt(ONE_PAGE_MARGIN_FLOOR_PT)
    for section, original_top, original_bottom in snapshot:
        section.top_margin = max(floor, Length(round(original_top * ratio)))
        section.bottom_margin = max(floor, Length(round(original_bottom * ratio)))


def _set_paragraph_mark_size(paragraph, pt_size: "Pt") -> None:
    """Sets the font size of a paragraph's own end-of-paragraph mark (the
    'rPr' nested inside 'pPr') — this is what determines the line height of
    an otherwise-EMPTY paragraph (zero runs), which the run-only loop in
    _set_all_font_sizes can never reach since there are no runs to iterate.

    Why this matters: a template can have several trailing blank paragraphs
    after its closing table (leftover spacer formatting, e.g. CHA TI.docx
    has 6 empty paragraphs after the main table). Left untouched, each one
    keeps whatever size the template originally authored — confirmed via
    inspection to be a FIXED 11pt regardless of template — at every ladder
    step, even once the ladder has shrunk everything else down to 9pt or
    7pt. That's a real, sizeable chunk of vertical space (roughly half an
    inch across 6 paragraphs) that font-size shrinking silently never
    reclaims, which is exactly why CHA TI needed to shrink all the way to
    9pt to fit one page despite having visible slack once it got there —
    confirmed by direct reproduction: forcing 11pt showed the ENTIRE real
    content (table, declaration, signature, footer) fitting on page 1 with
    room to spare, while page 2 was a phantom page containing nothing but
    the repeating header logo and footer, i.e. purely this unclaimed
    trailing-paragraph space pushed those last few blank paragraphs over.
    """
    half_points = str(int(round(pt_size.pt * 2)))
    pPr = paragraph._p.get_or_add_pPr()
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)
    for tag in ('w:sz', 'w:szCs'):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
        el.set(qn('w:val'), half_points)


def _set_all_font_sizes(doc: "DocxDocument", pt_size: float, pin_special_text: bool = True, amount_chargeable_pt: float = None) -> None:
    """Uniformly overwrites every run's font size — labels, headers and data
    alike — per the client's "shrink everything together" preference, EXCEPT
    the document title (ONE_PAGE_TITLE_TEXTS -> ONE_PAGE_TITLE_PT) and the
    genuine bottom "HORIZON ENTERPRISE" name/address block
    (ONE_PAGE_COMPANY_NAME_PT / ONE_PAGE_ADDRESS_PT), which stay pinned
    regardless of how far the rest of the page has to shrink — UNLESS
    pin_special_text=False, in which case those get uniformly sized like
    everything else too. That override exists purely as enforce_one_page's
    absolute last resort: on rare, content-heavy documents even font size +
    row-height + margin reclaiming all the way to their respective floors
    still isn't enough (confirmed by direct reproduction — a heavy Packing
    List still had the pinned footer alone spill a line onto page 2 with
    every other lever already maxed out), and a client requirement of
    "always one page" has to take priority over the title/footer pin in that
    rare case — every document that fits with the pin honored never reaches
    this parameter at all. Bold/italic/underline and font family are
    untouched either way; only the size changes.

    Also sets each paragraph's own end-of-paragraph mark size (see
    _set_paragraph_mark_size) — needed so EMPTY paragraphs (no runs) shrink
    along the same ladder instead of silently keeping their original size.

    Two passes, in this order:
    1. Company-name/address pin — restricted to
       _iter_body_and_header_footer_paragraphs (NOT table cells; see that
       function's docstring for why the table-cell lookalikes must be
       excluded). Each pinned paragraph's underlying XML element id is
       recorded so pass 2 skips it.
    2. Title pin + default ladder size for everything else — uses
       _iter_all_paragraphs (which DOES include table cells), since the
       title itself legitimately lives inside a table cell in some
       templates (e.g. ANX C.docx).
    """
    size = Pt(pt_size)
    title_size = size if not pin_special_text else Pt(ONE_PAGE_TITLE_PT)
    company_size = size if not pin_special_text else Pt(ONE_PAGE_COMPANY_NAME_PT)
    address_size = size if not pin_special_text else Pt(ONE_PAGE_ADDRESS_PT)
    amount_chargeable_size = size if not pin_special_text else Pt(amount_chargeable_pt if amount_chargeable_pt is not None else ONE_PAGE_AMOUNT_CHARGEABLE_PT)

    # NOTE: track the actual `_p` XML elements themselves in this set (not
    # id(paragraph._p)) — the Paragraph wrapper objects returned by each
    # generator are transient, and once garbage-collected, Python is free to
    # reuse their id() for the next unrelated object created during pass 2's
    # iteration, causing false "already pinned" positives on paragraphs that
    # were never actually pinned. Keeping the real elements here keeps them
    # alive and makes membership checks reliable.
    pinned_elements = set()
    for paragraph in (_iter_body_and_header_footer_paragraphs(doc) if pin_special_text else []):
        text_upper = paragraph.text.strip().upper()
        if text_upper in ONE_PAGE_COMPANY_NAME_TEXTS:
            paragraph_fallback_size = company_size
        elif _is_footer_address_paragraph(text_upper):
            paragraph_fallback_size = address_size
        else:
            continue
        pinned_elements.add(paragraph._p)
        # Classify PER RUN, not uniformly for the whole paragraph: some
        # templates (e.g. new_commercial_invoice1.docx) put "HORIZON
        # ENTERPRISE" and the address in ONE paragraph joined by a manual
        # line break rather than two separate paragraphs — paragraph.text
        # then reads as "HORIZON ENTERPRISE\nE-802, Safal Parisar..." as one
        # combined string, which contains "SAFAL PARISAR" and so classified
        # the ENTIRE paragraph (company name included) as address-sized,
        # which is the bug this fixes. Each run's OWN (line-break-stripped)
        # text is checked independently first; a run that doesn't
        # self-identify either way (e.g. a mid-address run split awkwardly
        # across multiple runs) falls back to the whole-paragraph
        # classification computed above, which is exactly the old,
        # already-correct behavior for the single-line-per-paragraph
        # templates (CHA TI/PL, Packing_List, PI FORMAT, ANX C).
        for run in paragraph.runs:
            run_text_upper = run.text.strip().upper()
            if run_text_upper in ONE_PAGE_COMPANY_NAME_TEXTS:
                run_size = company_size
            elif _is_footer_address_paragraph(run_text_upper):
                run_size = address_size
            else:
                run_size = paragraph_fallback_size
            run.font.size = run_size
        _set_paragraph_mark_size(paragraph, paragraph_fallback_size)

    for paragraph in _iter_all_paragraphs(doc):
        if paragraph._p in pinned_elements:
            continue
        text_upper = paragraph.text.strip().upper()
        if text_upper in ONE_PAGE_TITLE_TEXTS:
            run_size = title_size
        elif _is_amount_chargeable_paragraph(text_upper):
            run_size = amount_chargeable_size
        else:
            run_size = size
        for run in paragraph.runs:
            run.font.size = run_size
        _set_paragraph_mark_size(paragraph, run_size)


def _convert_to_pdf(docx_path: Path, out_dir: Path) -> Path:
    """Headless LibreOffice docx->pdf conversion, isolated to its own scratch
    user profile per call — soffice locks its profile directory, so two
    conversions running with the default profile at the same time can hang
    or fail. Raises on non-zero exit or missing output (caller decides how
    to handle: see enforce_one_page_).
    """
    profile_dir = out_dir / f"soffice_profile_{uuid.uuid4().hex}"
    profile_dir.mkdir(parents=True, exist_ok=True)
    try:
        subprocess.run(
            [
                "soffice", "--headless", "--norestore",
                f"-env:UserInstallation=file://{profile_dir}",
                "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path),
            ],
            check=True, capture_output=True, timeout=ONE_PAGE_SOFFICE_TIMEOUT_SEC,
        )
    finally:
        shutil.rmtree(profile_dir, ignore_errors=True)
    pdf_path = out_dir / (docx_path.stem + ".pdf")
    if not pdf_path.exists():
        raise RuntimeError(f"soffice did not produce {pdf_path}")
    return pdf_path


def _count_pdf_pages(pdf_path: Path) -> int:
    result = subprocess.run(
        ["pdfinfo", str(pdf_path)], check=True, capture_output=True, text=True,
        timeout=ONE_PAGE_SOFFICE_TIMEOUT_SEC,
    )
    for line in result.stdout.splitlines():
        if line.startswith("Pages:"):
            return int(line.split(":", 1)[1].strip())
    raise RuntimeError(f"pdfinfo output for {pdf_path} had no 'Pages:' line")


def _amount_chargeable_is_one_line(pdf_path: Path) -> bool:
    """True if the rendered "AMOUNT CHARGEABLE..." line fits on a single
    visual line rather than wrapping — checked for real via the actual PDF
    text layout (pdftotext -layout), not guessed. The sentence always ends
    in "...Only" (see convert_to_words / script.gs's AMOUNTWORDS), so a line
    that starts with "AMOUNT CHARGEABLE" and also ends with "Only" hasn't
    wrapped; if it wrapped, that first visual line would end mid-sentence
    and "Only" would appear at the start of the NEXT line instead."""
    result = subprocess.run(
        ["pdftotext", "-layout", str(pdf_path), "-"], check=True, capture_output=True, text=True,
        timeout=ONE_PAGE_SOFFICE_TIMEOUT_SEC,
    )
    for line in result.stdout.splitlines():
        stripped = line.strip()
        if stripped.upper().startswith("AMOUNT CHARGEABLE"):
            return stripped.rstrip(".").split()[-1].upper() == "ONLY"
    return True  # line not found at all (template variant without it) — nothing to fix


def _fit_amount_chargeable_one_line(doc: "DocxDocument", docx_path: Path, tmp_dir: Path, template_name: str, body_size_pt: float) -> None:
    """Runs once, after the main page-fit ladder has already settled on one
    page at body_size_pt — a small dedicated search that shrinks ONLY the
    "AMOUNT CHARGEABLE..." pin (starting at ONE_PAGE_AMOUNT_CHARGEABLE_PT)
    down in 0.5pt steps until that line fits on one visual line instead of
    wrapping, verified for real via the rendered PDF at every step, same
    rigor as the main page-fit ladder. body_size_pt (whatever size the main
    ladder actually settled on) is re-applied on every step so the rest of
    the page stays exactly as the main ladder left it — this never mutates
    the module-level ONE_PAGE_AMOUNT_CHARGEABLE_PT default, only what's
    written into this one document. Gives up at ONE_PAGE_MIN_PT (same floor
    as the main ladder) if it never manages to fit, leaving the line at
    whatever size it last tried.
    """
    for size_pt in _one_page_size_ladder(ONE_PAGE_AMOUNT_CHARGEABLE_PT, ONE_PAGE_MIN_PT):
        _set_all_font_sizes(doc, body_size_pt, amount_chargeable_pt=size_pt)
        doc.save(str(docx_path))

        pdf_path = _convert_to_pdf(docx_path, tmp_dir)
        one_line = _amount_chargeable_is_one_line(pdf_path)
        pdf_path.unlink(missing_ok=True)

        logger.info(f"[amount-chargeable] {template_name} @ {size_pt}pt -> {'one line' if one_line else 'wrapped'}")
        if one_line:
            return


def enforce_one_page(docx_path: Path, template_name: str) -> None:
    """Mutates the saved .docx in place so it renders to one page, per the
    client's request — skipped entirely for DBK (see
    ONE_PAGE_EXCLUDED_TEMPLATES; that document is expected to run long).
    Any failure here (soffice missing, conversion error, etc.) is logged and
    swallowed rather than failing the whole generation — the document still
    exists at whatever size was last successfully set, just not guaranteed
    one-page, which is strictly better than losing the document entirely.
    """
    if template_name in ONE_PAGE_EXCLUDED_TEMPLATES:
        return

    floor_pt = ONE_PAGE_MIN_PT
    ceiling_pt = ONE_PAGE_TEMPLATE_MAX_PT_OVERRIDE.get(template_name, ONE_PAGE_DEFAULT_PT)
    size_ladder = _one_page_size_ladder(ceiling_pt, floor_pt)

    with tempfile.TemporaryDirectory(prefix="one_page_check_") as tmp_dir_str:
        tmp_dir = Path(tmp_dir_str)
        try:
            doc = DocxDocument(str(docx_path))
            row_height_snapshot = _snapshot_row_heights(doc)
            margin_snapshot = _snapshot_margins(doc)
            for size_pt in size_ladder:
                _set_all_font_sizes(doc, size_pt)
                ratio = ONE_PAGE_ROW_HEIGHT_BASELINE_RATIO * (size_pt / ONE_PAGE_DEFAULT_PT)
                _scale_row_heights(row_height_snapshot, ratio)
                _scale_margins(margin_snapshot, ratio)
                doc.save(str(docx_path))

                pdf_path = _convert_to_pdf(docx_path, tmp_dir)
                pages = _count_pdf_pages(pdf_path)
                pdf_path.unlink(missing_ok=True)

                logger.info(f"[one-page] {template_name} @ {size_pt}pt -> {pages} page(s)")
                if pages <= 1:
                    _fit_amount_chargeable_one_line(doc, docx_path, tmp_dir, template_name, size_pt)
                    return

            # Absolute last resort: the whole ladder ran at the floor font
            # size with row heights and margins both fully reclaimed, and it
            # STILL doesn't fit — confirmed by direct reproduction to happen
            # on genuinely heavy documents (many models/vehicles) where even
            # 6pt body text isn't enough, because the pinned title/company/
            # address block doesn't shrink with the rest of the page. "Always
            # one page" wins over the pin at this point — try once more at
            # the floor size with the pin itself lifted.
            logger.warning(
                f"[one-page] {template_name} still {pages} page(s) at the {floor_pt}pt "
                "floor with row heights/margins fully reclaimed — trying one "
                "final pass with the title/company/address size pin lifted."
            )
            _set_all_font_sizes(doc, floor_pt, pin_special_text=False)
            doc.save(str(docx_path))
            pdf_path = _convert_to_pdf(docx_path, tmp_dir)
            pages = _count_pdf_pages(pdf_path)
            pdf_path.unlink(missing_ok=True)
            logger.info(f"[one-page] {template_name} @ {floor_pt}pt (pin lifted) -> {pages} page(s)")
            if pages <= 1:
                return
            logger.error(
                f"[one-page] {template_name} still {pages} page(s) even with every lever "
                "(font size, row heights, margins, title/footer pin) fully exhausted — "
                "leaving as-is. This document's real content genuinely exceeds one page "
                f"at {floor_pt}pt; only trimming actual content can fix it further."
            )
        except Exception as e:
            logger.error(f"[one-page] {template_name} enforcement failed, leaving document as rendered: {e}", exc_info=True)

# Maps each DOCX template to the description field it should use from Item
TEMPLATE_DESC_FIELD = {
    "PI FORMAT.docx":              "description_pi",
    "Commercial_Invoice.docx":     "description_commercial",
    "new_commercial_invoice1.docx": "description_commercial",
    "Packing_List.docx":           "description_packing",
    "Tax_Invoice.docx":            "description_tax",
    "SCOMET_Declaration.docx":     "description_scomet",
    "Annexure_C.docx":             "description_scomet",
    "ANX C.docx":                  "description_scomet",
    "Annexure_1.docx":             "description_annexure1",
    "CHA TI.docx":                 "description_cha_ti",
    "CHA PL.docx":                 "description_cha_pl",
    "CHA CI.docx":                 "description_cha_ci",
}

# Maps each DOCX template to which per-item HSN field it should print.
# Customer-facing documents (CI, PI, Tax Invoice, Packing List) use the
# customer's own country's HSN code; everything else — SCOMET/Annexure/DBK/
# Vintage and every CHA document — uses the Indian HSN code, since those are
# Indian customs/export filings. See the HSN resolution block in
# buildPayload() (script.gs).
TEMPLATE_HSN_FIELD = {
    "PI FORMAT.docx":              "hsn_code_pi",
    "Commercial_Invoice.docx":     "hsn_code_user_country",
    "new_commercial_invoice1.docx": "hsn_code_user_country",
    "Tax_Invoice.docx":            "hsn_code_user_country",
    "Packing_List.docx":           "hsn_code_user_country",
    "SCOMET_Declaration.docx":     "hsn_code_india",
    "Annexure_C.docx":             "hsn_code_india",
    "ANX C.docx":                  "hsn_code_india",
    "Annexure_1.docx":             "hsn_code_india",
    "DBK_Declaration.docx":        "hsn_code_india",
    "Vintage_Car_Declaration.docx":"hsn_code_india",
    "CHA TI.docx":                 "hsn_code_india",
    "CHA PL.docx":                 "hsn_code_india",
    "CHA CI.docx":                 "hsn_code_india",
}

# "Normal" Commercial Invoice + Packing List print the shipment's original
# PI invoice number + PI generation date as the Buyer's Order No. instead of
# whatever's typed in CONTROL!C15 (e.g. "HE/2026-27/PI/060 DT. 01.07.2026").
# Deliberately does NOT include CHA TI/PL/CI — those templates also
# reference {{ buyers_order_no }} but keep printing the CONTROL!C15 value
# as before; see the override block in build_context() below.
BUYERS_ORDER_NO_PI_COMBO_TEMPLATES = {"new_commercial_invoice1.docx", "Packing_List.docx"}

def sanitize_filename(name: str) -> str:
    return name.replace("/", "-").replace("\\", "-").replace(" ", "_")

# Monetary + rate fields — exchange_rate keeps its decimals (genuine currency
# conversion precision), everything else here is always a whole number in
# practice (GST slabs are 0/5/12/18/28) so it gets rounded to an int so
# templates never print "1500.0" or "18.0%".
#
# USD amounts use Western thousands-grouping (1,850 / 70,750) — that's the
# correct convention for a $ figure on an export document. INR amounts must
# use Indian digit grouping instead (last 3 digits together, then pairs of 2
# — e.g. 6650500 -> "66,50,500", not the Western "6,650,500" they were
# printing before) — every CHA/CI/PL/Tax Invoice template that shows a ₹
# figure already has the ₹ symbol baked in as static template text right
# before the placeholder, so the formatter itself only needs to produce the
# grouped digits.
_MONEY_KEYS_USD = (
    'unit_price_usd', 'fob_total_usd', 'freight_usd', 'insurance_usd', 'cif_total_usd',
    'total_cif_usd', 'total_fob_usd',
)
_MONEY_KEYS_INR = (
    'taxable_value_inr', 'igst_amount_inr', 'total_value_inr', 'total_inr', 'total_taxable_inr',
)
_MONEY_KEYS_OTHER = ('igst_rate', 'igst_percent')  # percentages, not currency — never need grouping
_MONEY_KEYS_ITEM = ('rate_per_unit', 'amount_usd', 'unit_price', 'total')  # always USD
_MONEY_KEYS_VEHICLE = ('unit_price_usd',)  # always USD

# Genuine rates/measurements — real decimals (e.g. 83.45 exchange rate, 600.5
# kg) must survive untouched, but a value that happens to be a whole number
# (e.g. 94.0, 3600.0) should still print as "94" / "3600", not "94.0" /
# "3600.0". So: strip the ".0" only when there's no actual fraction.
_RATE_KEYS_TOP = (
    'exchange_rate', 'total_net_weight', 'total_gross_weight', 'net_weight_kg', 'gross_weight_kg'
)

def _clean_rate(v):
    return int(v) if isinstance(v, float) and v == int(v) else v

def _format_money(v):
    """Whole-number USD amount as a Western thousands-comma string (1850 ->
    '1,850'). Non-numeric values (already a string, None, etc.) pass through
    unchanged."""
    if isinstance(v, bool):
        return v
    if isinstance(v, (int, float)):
        return f"{round(v):,}"
    return v

def _format_money_indian(v):
    """Whole-number INR amount using Indian digit grouping — last 3 digits
    together, then pairs of 2 from there (6650500 -> '66,50,500'). Mirrors
    the client-approved formatIndianPrice() grouping rule. Non-numeric
    values pass through unchanged, same as _format_money."""
    if isinstance(v, bool):
        return v
    if isinstance(v, (int, float)):
        n = round(v)
        digits = str(abs(n))
        if len(digits) <= 3:
            grouped = digits
        else:
            last_three, rest = digits[-3:], digits[:-3]
            pairs = []
            while len(rest) > 2:
                pairs.insert(0, rest[-2:])
                rest = rest[:-2]
            if rest:
                pairs.insert(0, rest)
            grouped = ','.join(pairs) + ',' + last_three
        return ('-' if n < 0 else '') + grouped
    return v

def _round_money(context: Dict[str, Any]) -> None:
    for key in _MONEY_KEYS_USD:
        if key in context:
            context[key] = _format_money(context[key])
    for key in _MONEY_KEYS_INR:
        if key in context:
            context[key] = _format_money_indian(context[key])
    for key in _MONEY_KEYS_OTHER:
        if key in context:
            context[key] = _format_money(context[key])
    for key in _RATE_KEYS_TOP:
        if key in context:
            context[key] = _clean_rate(context[key])
    for it in (context.get('items') or []):
        if isinstance(it, dict):
            for key in _MONEY_KEYS_ITEM:
                if key in it:
                    it[key] = _format_money(it[key])
    item_singular = context.get('item')
    if isinstance(item_singular, dict):
        for key in _MONEY_KEYS_ITEM:
            if key in item_singular:
                item_singular[key] = _format_money(item_singular[key])
    for v in (context.get('vehicles') or []):
        if isinstance(v, dict):
            for key in _MONEY_KEYS_VEHICLE:
                if key in v:
                    v[key] = _format_money(v[key])

def convert_to_words(amount: float, currency: str = "INR") -> str:
    try:
        major = int(amount)
        minor = int(round((amount - major) * 100))
        if currency == "INR":
            major_words = num2words(major, lang='en_IN').title()
            currency_name = "Rupees"
        else:
            major_words = num2words(major, lang='en').title()
            currency_name = "US Dollars"
        word_str = f"{major_words} {currency_name}"
        if minor > 0:
            minor_words = num2words(minor, lang='en_IN' if currency == "INR" else 'en').title()
            word_str += f" And {minor_words} {'Paise' if currency == 'INR' else 'Cents'}"
        return word_str + " Only"
    except Exception as e:
        logger.error(f"Words conversion error: {e}")
        return f"{amount} {currency}"

def build_context(payload: Any, template_name: str = "") -> Dict[str, Any]:
    if hasattr(payload, 'model_dump'):
        context = payload.model_dump()
    else:
        context = dict(payload) if isinstance(payload, dict) else payload

    # ── Flatten financials ──────────────────────────────────────────────────
    if isinstance(context.get('financials'), dict):
        financials = context.pop('financials')
        context.update({
            'quantity':          financials.get('quantity', 0),
            'unit_price_usd':    financials.get('unit_price_usd', 0),
            'fob_total_usd':     financials.get('fob_total_usd', 0),
            'freight_usd':       financials.get('freight_usd', 0),
            'insurance_usd':     financials.get('insurance_usd', 0),
            'cif_total_usd':     financials.get('cif_total_usd', 0),
            'exchange_rate':     financials.get('exchange_rate', 0),
            'igst_rate':         financials.get('igst_rate', 0),
            'taxable_value_inr': financials.get('taxable_value_inr', 0),
            'igst_amount_inr':   financials.get('igst_amount_inr', 0),
            'total_value_inr':   financials.get('total_value_inr', 0),
            'total_inr':         financials.get('total_value_inr', 0),
        })

    # ── Freight/insurance boolean flags ─────────────────────────────────────
    # Computed here (numeric, pre-formatting) rather than letting templates
    # test freight_usd/insurance_usd directly — _round_money() below turns
    # those into thousands-comma STRINGS (e.g. 0 -> "0"), and a non-empty
    # string is truthy in Jinja even when its value is "0". A template-side
    # `{% if freight_usd %}` would therefore always be true. These two flags
    # stay real booleans all the way to render, so CHA CI's freight/insurance
    # rows and the CI/CHA TI "CIF vs FOB" label can safely branch on them.
    context['has_freight']              = bool(context.get('freight_usd'))
    context['has_insurance']            = bool(context.get('insurance_usd'))
    context['has_freight_or_insurance'] = context['has_freight'] or context['has_insurance']

    # ── Flatten weights ─────────────────────────────────────────────────────
    if isinstance(context.get('weights'), dict):
        weights = context.get('weights', {})
        context['total_packages']     = weights.get('total_packages', 0)
        context['total_net_weight']   = weights.get('net_weight_kg', 0)
        context['total_gross_weight'] = weights.get('gross_weight_kg', 0)
        context['net_weight_kg']      = weights.get('net_weight_kg', 0)
        context['gross_weight_kg']    = weights.get('gross_weight_kg', 0)

    # ── Flatten shipping (kept as nested dict too for PI FORMAT shipping.* refs) ──
    if isinstance(context.get('shipping'), dict):
        shipping = context.get('shipping', {})
        context['container_no']           = shipping.get('container_no', '')
        context['container_type']         = shipping.get('container_type', '')
        context['pre_carriage_by']        = shipping.get('pre_carriage_by', '')
        context['mode_of_transport']      = shipping.get('mode_of_transport', '')
        context['country_of_origin']      = shipping.get('country_of_origin', 'INDIA')
        context['port_of_loading']        = shipping.get('port_of_loading', '')
        context['port_of_discharge']      = shipping.get('port_of_discharge', '')
        context['final_destination']      = shipping.get('final_destination', '')
        context['country_of_destination'] = shipping.get('country_of_destination', '')
        context['place_of_receipt']  = context.get('place_of_receipt')  or shipping.get('port_of_loading', '')
        context['delivery_terms']    = context.get('delivery_terms')    or shipping.get('mode_of_transport', '')
    else:
        context['container_no'] = context.get('container_no', '')

    # ── Cross-template name aliases ─────────────────────────────────────────
    context['total_cif_usd']     = context.get('cif_total_usd') or 0
    context['total_fob_usd']     = context.get('total_fob_usd') or context.get('fob_total_usd') or 0
    context['total_taxable_inr'] = context.get('taxable_value_inr') or 0
    context['igst_percent']      = context.get('igst_rate') or 0

    # ── Swap item.description per template ──────────────────────────────────
    # Always assign (not "only if truthy") — an intentionally blank desc_field
    # (Invoice_Descriptions collapsing every item but the first down to '' so
    # a concatenating template prints the text exactly once) must make
    # `description` blank too, not silently fall back to whatever value it
    # already held from an earlier per-template pass.
    desc_field = TEMPLATE_DESC_FIELD.get(template_name, "")
    if desc_field and isinstance(context.get('items'), list):
        for it in context['items']:
            if isinstance(it, dict):
                # Always set 'description', even if the source field is missing (use fallback to description_pi)
                if desc_field in it:
                    it['description'] = it[desc_field]
                elif 'description_pi' in it:
                    # Fallback to description_pi if the template-specific field doesn't exist
                    it['description'] = it['description_pi']
                else:
                    # Last resort: use empty string
                    it['description'] = ''

    # ── Swap item.hsn_code per template ──────────────────────────────────────
    # hsn_code_user_country feeds Commercial_Invoice/Tax_Invoice/Packing_List,
    # hsn_code_india feeds everything else (SCOMET/Annexure/CHA), hsn_code_pi
    # feeds PI FORMAT. Each of those 3 fields already falls back to the
    # existing model-level HSN in buildPayload() (script.gs) when a Stock tab
    # override is blank, so this only ever changes output once an override
    # is actually filled in. Runs before context['item'] (singular, below) is
    # copied from items[0], so that also picks up the already-swapped value.
    hsn_field = TEMPLATE_HSN_FIELD.get(template_name, "")
    if hsn_field and isinstance(context.get('items'), list):
        for it in context['items']:
            if isinstance(it, dict) and it.get(hsn_field):
                it['hsn_code'] = it[hsn_field]

    # ── Buyer's Order No. override — PI number + PI date (CI/PL only) ──────
    # See BUYERS_ORDER_NO_PI_COMBO_TEMPLATES above. Resolved from vehicles[],
    # same source (and same "first vehicle with a non-blank pi_invoice_no
    # wins" rule) as script.gs's buildPayload() uses for its own resolved-PI
    # values — every vehicle on a shipment was assigned together, so the
    # first non-blank hit is authoritative. Falls back to leaving
    # buyers_order_no untouched (whatever CONTROL!C15 sent) if this shipment
    # never had a PROFORMA stage recorded on any vehicle.
    if template_name in BUYERS_ORDER_NO_PI_COMBO_TEMPLATES and isinstance(context.get('vehicles'), list):
        pi_invoice_no = ''
        pi_invoice_date = ''
        for v in context['vehicles']:
            if isinstance(v, dict) and v.get('pi_invoice_no'):
                pi_invoice_no = v['pi_invoice_no']
                pi_invoice_date = v.get('pi_invoice_date', '')
                break
        if pi_invoice_no:
            # No "DT." label — just the PI number, two spaces, then the date
            # (e.g. "HE/2026-27/PI/060  09.08.2026"), per client request.
            context['buyers_order_no'] = f"{pi_invoice_no}  {pi_invoice_date}" if pi_invoice_date else pi_invoice_no

    # ── Add aliases + sr_no + unit + package range on each item ────────────
    _running_item = 0
    for idx, it in enumerate(context.get('items') or []):
        if isinstance(it, dict):
            it['sr_no']   = idx + 1
            it.setdefault('unit',       'Nos')
            it.setdefault('unit_price', it.get('rate_per_unit', 0))
            it.setdefault('total',      it.get('amount_usd', 0))
            _qty = it.get('quantity', 0)
            it['sr_start'] = _running_item + 1
            it['sr_end']   = _running_item + _qty
            _running_item += _qty
            # Ensure all description fields exist (fallback to empty string if missing)
            for desc_key in ['description', 'description_commercial', 'description_scomet',
                            'description_packing', 'description_tax', 'description_pi',
                            'description_annexure1', 'description_cha_ti', 'description_cha_pl', 'description_cha_ci']:
                it.setdefault(desc_key, '')
            # Ensure HSN fields exist
            for hsn_key in ['hsn_code', 'hsn_code_pi', 'hsn_code_india', 'hsn_code_user_country']:
                it.setdefault(hsn_key, '')

    # ── total_quantity across all items ─────────────────────────────────────
    context['total_quantity'] = sum(
        it.get('quantity', 0) for it in (context.get('items') or [])
        if isinstance(it, dict)
    )
    # PI FORMAT uses top-level {{ quantity }} — override with actual vehicle count
    if context['total_quantity'] > 0:
        context['quantity'] = context['total_quantity']

    # ── item singular for PI FORMAT — aggregate all items into one summary row
    if isinstance(context.get('items'), list) and context['items']:
        first = dict(context['items'][0])
        first['quantity']   = context['total_quantity']
        first['amount_usd'] = sum(it.get('amount_usd', 0) for it in context['items'] if isinstance(it, dict))
        first['total']      = first['amount_usd']
        context['item'] = first

    # PI FORMAT.docx and Tax_Invoice.docx both use a native docxtpl row-loop
    # ({%tr for item in items %}, referencing item.hsn_code/description/
    # quantity/rate_per_unit/amount_usd — Tax_Invoice also uses item.sr_start/
    # sr_end) — they were already built to print one row per model, same as
    # Commercial Invoice / CHA CI/TI/PL. A previous version of this function
    # collapsed context['items'] into one merged summary row here, which
    # fought the template's own design and silently dropped every model but
    # the first from the description. Removed — context['items'] (with each
    # item's sr_no/sr_start/sr_end already set above) now renders as-is, one
    # row per model, exactly like the other multi-model documents.

    # ── vin_list for Annexure_1 ({% for v in vin_list %}) ──────────────────
    if isinstance(context.get('vehicles'), list):
        context['vin_list'] = [
            {
                'sr_no':      i + 1,
                'chassis_no': v.get('chassis_no', ''),
                'engine_no':  v.get('engine_no', ''),
                'model':      v.get('model', ''),
                'color':      v.get('color', ''),
            }
            for i, v in enumerate(context['vehicles'])
            if isinstance(v, dict)
        ]
        # vehicle_models_list for Annexure C (item 15 "Vehicles" section) —
        # script.gs resolves this per-invoice (auto-filled once from
        # Invoice_Descriptions!L, then sticks even if hand-edited there) and
        # sends it as annexure_c_vehicles_list; this only computes a fallback
        # (distinct models, comma-separated) for direct-API calls that don't
        # go through script.gs at all.
        context['vehicle_models_list'] = context.get('annexure_c_vehicles_list') or ', '.join(dict.fromkeys(
            v.get('model', '') for v in context['vehicles']
            if isinstance(v, dict) and v.get('model')
        ))

    # ── notify_1 fallback → buyer name + address for Annexure C consignee ──
    if not context.get('notify_1'):
        buyer = context.get('buyer') or {}
        if isinstance(buyer, dict):
            parts = [buyer.get('name', ''), buyer.get('address', ''), buyer.get('country', '')]
            context['notify_1'] = ', '.join(p for p in parts if p)

    # ── packages for Packing_List ({% for p in packages %}) ────────────────
    if isinstance(context.get('items'), list):
        marks = context.get('marks_and_numbers', '')
        _running_pkg = 0
        _pkg_list = []
        for it in context['items']:
            if isinstance(it, dict):
                _qty = it.get('quantity', 0)
                _pkg_list.append({
                    'marks':       marks,
                    'description': it.get('description', ''),
                    'hsn_code':    it.get('hsn_code', ''),
                    'quantity':    _qty,
                    'unit':        it.get('unit', 'Nos'),
                    'sr_start':    _running_pkg + 1,
                    'sr_end':      _running_pkg + _qty,
                })
                _running_pkg += _qty
        context['packages'] = _pkg_list

    # ── generation_date fallback ────────────────────────────────────────────
    if not context.get('generation_date'):
        context['generation_date'] = datetime.utcnow().strftime('%d.%m.%Y')

    # ── scomet_product_desc fallback ────────────────────────────────────────
    # Joins every item's description (not just the first) so multi-model
    # shipments list all products, comma-separated, in the declaration letter.
    if not context.get('scomet_product_desc') and isinstance(context.get('items'), list) and context['items']:
        descs = [
            (it.get('description_scomet') or it.get('description') or '')
            for it in context['items'] if isinstance(it, dict)
        ]
        context['scomet_product_desc'] = ', '.join(d for d in descs if d)

    # ── insurance_ref_no fallback ───────────────────────────────────────────
    if not context.get('insurance_ref_no'):
        context['insurance_ref_no'] = context.get('lc_number', '')

    # ── container_type fallback (Annexure C item 14) ────────────────────────
    # Fed from CONTROL!F37 via script.gs shipping.container_type. Falls back
    # to the old hardcoded value when that cell is left blank.
    if not context.get('container_type'):
        context['container_type'] = "40' HQ"

    # ── Amount in words ─────────────────────────────────────────────────────
    if not context.get('amount_usd_words'):
        context['amount_usd_words'] = convert_to_words(context.get('cif_total_usd') or 0, currency="USD")
    if not context.get('amount_inr_words'):
        context['amount_inr_words'] = convert_to_words(
            context.get('total_value_inr') or context.get('total_inr') or 0, currency="INR"
        )

    # PI FORMAT's amount cell is a bare {{ amount_usd_words }} with no static
    # label (unlike CI/CHA CI/CHA TI, which already print their own "AMOUNT
    # CHARGEABLE..." label before the placeholder) — so PI FORMAT needs the
    # full phrase baked into the value itself.
    if template_name == "PI FORMAT.docx":
        for key in ("amount_usd_words", "amount_inr_words"):
            val = context.get(key) or ""
            if val and not val.upper().startswith("AMOUNT CHARGEABLE"):
                context[key] = f"AMOUNT CHARGEABLE IN {val}"

    # ── Whole-number money — strip ".0" from every USD/INR amount ──────────
    _round_money(context)

    context['generation_timestamp'] = datetime.utcnow().isoformat()
    logger.info(f"[{template_name}] Context built: {len(context)} keys | invoice={context.get('invoice_no')} | items={len(context.get('items') or [])} | vehicles={len(context.get('vehicles') or [])}")
    return context

def generate_document(template_name: str, payload: Any, invoice_no: str) -> Dict[str, Any]:
    safe_invoice = sanitize_filename(invoice_no)
    logger.info(f"[{invoice_no}] Generating {template_name}")

    try:
        template_path = TEMPLATE_DIR / template_name
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")

        context = build_context(payload, template_name=template_name)

        logger.info(f"Rendering template {template_name} with {len(context)} context keys")
        logger.info(f"Items in context: {len(context.get('items', []))} | Vehicles: {len(context.get('vehicles', []))}")
        if context.get('items'):
            for idx, item in enumerate(context['items'][:2]):  # Log first 2 items for debugging
                if isinstance(item, dict):
                    logger.info(f"  Item[{idx}]: hsn={item.get('hsn_code')} qty={item.get('quantity')} desc_comm={bool(item.get('description_commercial'))}")

        doc = DocxTemplate(str(template_path))
        doc.render(context) 
        
        output_filename = f"{safe_invoice}_{template_name}"
        output_path = OUTPUT_DIR / output_filename
        doc.save(str(output_path))

        enforce_one_page(output_path, template_name)

        logger.info(f"[{invoice_no}] {template_name} generated locally at {output_path}")
        return {
            "status": "success",
            "template": template_name,
            "output_file": str(output_path),
            "download_name": output_filename,
            "download_url": f"{BASE_DOWNLOAD_URL}/{output_filename}",
        }
  
    except Exception as e:
        logger.error(f"[{invoice_no}] {template_name} FAILED: {str(e)}", exc_info=True)
        return {
            "status": "error",
            "template": template_name,
            "error": str(e)
        }
